"""
Anonimizador de textos juridicos brasileiros - LGPD Compliance v4.
Le .txt, .docx e .pdf de /entrada, grava anonimizados em /saida,
relatorios JSON em /relatorios e mapas de reversao em /cofre.

Tokens numerados por categoria: [NOME_01], [DOCUMENTO_01], etc.
Sem dependencias externas obrigatorias; .docx requer python-docx, .pdf requer pymupdf.

Conformidade LGPD (art. 37 e art. 6, VIII):
  - Cada relatorio inclui hash SHA-256 do original e metadados de tratamento.
  - Log consolidado em relatorios/log_compliance.jsonl (append-only).
  - Metadados configurados em config.json (criado automaticamente se ausente).
  - Mapa de reversao (cofre) salvo em cofre/<nome>_mapa.json (pseudonimizacao reversivel).

NOVO em v4 (reescrita do motor apos falhas graves detectadas na v3.1):
  - Motor baseado em SPANS: todos os padroes sao detectados sobre o texto ORIGINAL,
    as sobreposicoes sao resolvidas (span mais longo vence) e as substituicoes sao
    aplicadas em uma unica passada. Elimina o hack de protecao (paragrafos engolidos,
    texto colado ao token, dupla substituicao).
  - DOCX: processa TODOS os paragrafos, TABELAS (inclusive aninhadas), cabecalhos e
    rodapes. A v3 ignorava tabelas por completo.
  - DOCX: substituicao por offset dentro dos runs — preserva formatacao e NUNCA
    desalinha paragrafos (a v3 aplicava substituicoes nas linhas erradas).
  - Consistencia de entidade: o mesmo dado recebe o MESMO token em todo o documento.
  - Propagacao: um nome detectado uma vez (ex.: qualificacao do socio) e substituido
    em TODAS as demais ocorrencias, inclusive em tabelas e assinaturas, com e sem
    sufixo societario (PRIMAVERA SABORES LTDA -> tambem "Primavera Sabores").
  - Novos padroes: RG por gatilho (carteira de identidade n. 134.226), NIRE,
    valores com extenso (R$ 30.700,00 (trinta mil e setecentos reais)),
    quantidade de quotas com extenso, enderecos completos ate o CEP, cidade/UF.
  - Correcao: razoes sociais com acentos (PRIMAVERA PARTICIPACOES E INVESTIMENTOS LTDA).
  - Correcao: removido o padrao "sociedade|empresa + qualquer texto" que gerava
    tokens-lixo com a palavra generica "Sociedade" e engolia texto adjacente.

NOVO em v4.1:
  - Saida de PDF em .txt ou .docx (CLI: --saida-pdf; app: seletor "Saida do PDF").
  - Nomes com inicial abreviada (Fernando M. Sampaio, Manoel M. da Silva Prado Neto).
  - Sociedades de advogados (SAMPAIO E CAMPOS ADVOGADOS, Silva Advocacia), com guardas
    para nao capturar 'Ordem dos Advogados' e 'Estatuto da Advocacia'.
  - Propagacao de SOBRENOMES: nome detectado propaga tambem as duas ultimas
    palavras fortes ('Dr. Sampaio', 'Silva Prado'); sufixo de geracao (Neto,
    Filho, Junior) nunca propaga isolado.
  - Correcao: abreviacoes de logradouro (R., Av.) agora exigem maiuscula e
    palavra capitalizada em seguida — 'a r. sentenca de fls.' nao e mais
    capturado como endereco (falso positivo grave que engolia texto).
  - EGREGIO/COLENDA/EXMO adicionados as stopwords de nome.

Uso:
  python anonimizador.py                        # anonimiza tudo em /entrada
  python anonimizador.py --saida-pdf txt        # PDFs geram saida .txt (ou docx|pdf)
  python anonimizador.py --restaurar <arquivo>  # restaura arquivo usando mapa em /cofre

Limitacao PDF: apenas PDFs digitais (texto selecionavel). PDFs escaneados (imagem)
exigem OCR e nao sao suportados nesta versao.
"""
import re, json, os, sys, unicodedata, hashlib, getpass, tempfile, shutil, bisect
from datetime import datetime

# ---------------------------------------------------------------------------
# Utilitarios de texto
# ---------------------------------------------------------------------------

def _sem_acento(s):
    """Remove acentos para comparacao com STOPWORDS (que usam ASCII)."""
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')


def _norm_entidade(dado):
    """Chave de normalizacao para deduplicar entidades (mesmo dado -> mesmo token)."""
    return _sem_acento(' '.join(dado.split())).casefold().strip(' .,;:')


def _cat_publica(categoria):
    """Categoria publicada nos tokens/relatorios. Variantes internas de NOME
    (NOME_ORG) sao publicadas simplesmente como NOME."""
    return 'NOME' if categoria.startswith('NOME') else categoria

# ---------------------------------------------------------------------------
# Dependencias opcionais
# ---------------------------------------------------------------------------

try:
    from docx import Document as _DocxDocument
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

try:
    import fitz as _fitz
    PDF_DISPONIVEL = True
except ImportError:
    PDF_DISPONIVEL = False

# ---------------------------------------------------------------------------
# Stopwords
# ---------------------------------------------------------------------------

STOPWORDS_NOME = {
    'ESTADO','GOVERNO','MUNICIPIO','RONDONIA','BRASIL','UNIAO',
    'FUNDACAO','HEMATOLOGIA','HEMOTERAPIA','ADMINISTRACAO',
    'ASSESSORIA','JURIDICA','TERMO','ADITIVO','CONTRATO','CONTRATANTE','CONTRATADA',
    'SECRETARIA','MINISTERIO','PREFEITURA','AUTARQUIA','TRIBUNAL','JUSTICA',
    'RUA','AVENIDA','PORTO','VELHO','RIO','JANEIRO','BENFICA','PRESIDENTE','EMPRESA',
    'SOCIO','SOCIA','SOCIOS','SOCIAS','SOCIEDADE','HOLDING','COMPANHIA',
    'ADMINISTRADOR','ADMINISTRADORA','REPRESENTANTE','LABORATORIOS','CONTROLE',
    'QUALIDADE','PARA','ASSEJUR','FHEMERON',
    'FEDERAL','ESTADUAL','MUNICIPAL','NACIONAL','DISTRITAL','REGIONAL','SUPERIOR','ESPECIAL',
    'CAMARA','SENADO','CONGRESSO','DEFENSORIA','PUBLICO',
    'PODER','EXECUTIVO','LEGISLATIVO','JUDICIARIO',
    'COORDENACAO','DIRETORIA','GERENCIA','SUBSECRETARIA',
    'PROCURADORIA','DELEGACIA','SUPERINTENDENCIA',
    'RECEITA','PGFN','CARF','CSRF','STF','STJ','TST','TSE','STM',
    # Termos processuais do eproc/TJ que nao sao nomes
    'CAPA','PROCESSO','ABERTURA','MOVIMENTO','SITUACAO','COMPETENCIA',
    'CLASSE','ASSUNTOS','CODIGO','DESCRICAO','PRINCIPAL','TUTELA',
    'URGENCIA','PROVISORIA','LIMINAR','ANTECIPACAO','INFORMACOES','ADICIONAIS',
    'CHAVE','NIVEL','SIGILO','ANEXOS','ELETRONICOS','CRIANCA','ADOLESCENTE',
    'DOENCA','GRAVE','DEVEDOR','IDOSO','DEFICIENCIA','PRIORIDADE','TRAMITACAO',
    'PESSOA','FISICA','JURIDICA','REPRESENTANTES','PARTES','VALOR','CAUSA',
    'GRATUITA','DIGITAL','SEPARACAO','PAGINA','EVENTO','INICIAL','FINAL',
    'PROCEDIMENTO','COMUM','CIVEL','CRIMINAL','FAZENDA','PUBLICA','VARA',
    'COMARCA','CARTORIO','ESCRIVANIA','DISTRIBUICAO','SORTEIO','SISTEMA',
    'USUARIO','OPERADOR','SERVIDOR','CERTIDAO','AUTENTICACAO','BANCARIA',
    'GRERJ','ITBI','ISS','IPTU','IPVA','ICMS','IPI','COFINS','CSLL','IRPJ',
    'MANIFESTA','REQUER','REQUEREMOS','INFORMA','SOLICITA','DECLARA','AFIRMA','ALEGA',
    'DESINTERESSE','AUDIENCIA','CONCILIACAO','MEDIACAO',
    'OPCAO','JUIZO','ARBITRAL','CONVENCIONAL',
    'COMPLEMENTARES','TRIBUTOS','RESERVADO','FISCO',
    'ACOMODACAO','INTERNACAO','PADRAO','AREA','ATUACAO','FORMACAO','PRECO',
    'AUTORIZACAO','PREVIA','DECLARACAO','SAUDE','CIVIL','TELEFONE',
    'RESIDENCIAL','CELULAR','COMERCIAL','ENDERECO','PROFISSIONAL',
    # Termos societarios/contratuais (v4)
    'CAPITAL','QUOTAS','COTAS','TOTAL','CLAUSULA','PARAGRAFO','DELIBERACAO',
    'DELIBERACOES','ADMISSAO','AUMENTO','CONFERENCIA','COMPOSICAO','RATIFICACAO',
    'CONSOLIDACAO','CONSOLIDADO','DENOMINACAO','DURACAO','SEDE','FORO','OBJETO',
    'SOCIAL','SOCIAIS','CESSAO','TRANSFERENCIA','REMUNERACAO','EXERCICIO',
    'RESULTADO','DISSOLUCAO','DECLARACOES','JUCERJA','JUCESP','JUCEMG','NIRE',
    'ALTERACAO','LIMITADA','UNIPESSOAL',
    # Pronomes de tratamento forense (nao sao nomes de pessoa)
    'EGREGIO','COLENDA','COLENDO','EXMO','EXMA','MERITISSIMO','MERITISSIMA',
}

# Palavras genericas que nunca sao, sozinhas, razao social
_GENERICOS_ORG = {'SOCIEDADE','EMPRESA','HOLDING','COMPANHIA','BANCO','INSTITUICAO',
                  # guardas para o padrao de sociedades de advogados: evita capturar
                  # 'Ordem dos Advogados', 'Estatuto da Advocacia', 'Exame ... Advocacia'
                  'ORDEM','ESTATUTO','COMISSAO','EXAME','ESCOLA'}

_MESES = r'(?:janeiro|fevereiro|mar[cç]o|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)'
_UFS = r'(?:AC|AL|AP|AM|BA|CE|DF|ES|GO|MA|MT|MS|MG|PA|PB|PR|PE|PI|RJ|RN|RS|RO|RR|SC|SP|SE|TO)'
# Abreviacoes de logradouro sao CASE-SENSITIVE e exigem palavra capitalizada
# (ou numero) na sequencia — evita capturar 'a r. sentenca' (respeitavel) e 'et al.'
_LOGRADOUROS = (r'(?:Rua|Avenida|Travessa|Alameda|Rodovia|Estrada|Pra[çc]a|Largo|'
                r'(?-i:(?:R|AV|Av|TRAV|Trav|AL|Al|ROD|Rod|EST|Est|LG|Lg)\.(?=[ \t]+[A-ZÀ-ÖØ-Þ0-9])))')
_SUFIXOS_SOC = r'(?:LTDA|Ltda|EIRELI|Eireli|EPP|ME|S\.A\.|S\.A|S/A|SLU)'
# Sufixos de sociedades de advogados e afins (sem sufixo empresarial classico)
_SUFIXOS_ADV = (r'(?:SOCIEDADE[ \t]+DE[ \t]+ADVOGADOS|Sociedade[ \t]+de[ \t]+Advogados|'
                r'ADVOGADOS[ \t]+ASSOCIADOS|Advogados[ \t]+Associados|'
                r'ADVOGADOS|Advogados|ADVOCACIA|Advocacia)')

# Palavras de nome proprio (com acentos)
_PAL_CAPS  = r'[A-ZÀ-ÖØ-Þ]{2,}'
_PAL_MISTA = r'[A-ZÀ-ÖØ-Þ][a-zà-öø-ÿ]{1,}'
_CONECT_CAPS = r'(?:D[AEO]S?|E)'
_CONECT_MIN  = r'(?:d[aeo]s?|e)'

# Inicial abreviada no meio do nome: Manoel M. da Silva Prado Neto
_INICIAL = r'[A-ZÀ-ÖØ-Þ]\.'
# Nome completo em caixa alta: PEDRO SOUZA SAMPAIO / ANA MARIA DE ANDRADE CAMPOS
_NOME_CAPS = _PAL_CAPS + r'(?:[ \t]+(?:' + _CONECT_CAPS + r'|' + _INICIAL + r'|' + _PAL_CAPS + r')){1,7}'
# Nome completo em caixa mista (aceita iniciais): Fernando M. Sampaio
_NOME_MISTO = _PAL_MISTA + r'(?:[ \t]+(?:' + _CONECT_MIN + r'|' + _INICIAL + r'|' + _PAL_MISTA + r')){1,7}'

# Qualificadores que seguem um nome de pessoa em pecas juridicas/contratos
_QUALIFICADORES = (
    r'(?:brasileir|portugu[eê]s|portuguesa|natural[ \t]+de|nacionalidade|'
    r'portador|portadora|inscrit[oa]|casad[oa]|solteir[oa]|divorciad[oa]|'
    r'vi[úu]v[oa]|separad[oa]|empres[áa]ri[oa]|advogad[oa]|m[ée]dic[oa]|'
    r'engenheir[oa]|comerciante|m[úu]sic[oa]|aposentad[oa]|servidor|funcion[áa]ri[oa]|'
    r'acima[ \t]+qualificad|j[áa][ \t]+qualificad|devidamente[ \t]+qualificad|'
    r'CPF|RG|CNPJ|OAB)'
)


# ---------------------------------------------------------------------------
# Validadores
# ---------------------------------------------------------------------------

def _nome_valido(nome):
    """Valida nome de pessoa: >=2 palavras fortes, sem stopword na primeira."""
    partes = re.findall(r'[A-Za-zÀ-ÖØ-öø-ÿ]+', nome)
    fortes = [p for p in partes if p.lower() not in {'de', 'da', 'do', 'dos', 'das', 'e'}]
    if len(fortes) < 2:
        return False
    if _sem_acento(fortes[0]).upper() in STOPWORDS_NOME:
        return False
    return not all(_sem_acento(p).upper() in STOPWORDS_NOME for p in fortes)


def _org_valida(nome):
    """Valida razao social: nao pode ser apenas palavra generica (Sociedade etc.)."""
    palavras = [p for p in re.findall(r'[A-Za-zÀ-ÖØ-öø-ÿ]+', nome)
                if p.lower() not in {'de', 'da', 'do', 'dos', 'das', 'e'}]
    if not palavras:
        return False
    if _sem_acento(palavras[0]).upper() in _GENERICOS_ORG:
        return False
    return True


def _num_nao_e_referencia_legal(texto, m):
    """Numero de milhar so vira VALOR se nao for artigo/inciso/numero de lei."""
    prefixo = texto[max(0, m.start() - 14):m.start()].lower()
    if re.search(r'(?:art\.?|artigo|inciso|lei[ \t]+n?[ºo°.]?|n[ºo°]\.?|§)\s*$', prefixo):
        return False
    if re.match(r'\s*/\s*\d', texto[m.end():m.end() + 8]):   # 9.249/1995
        return False
    return True


def _val_nome(texto, m):
    return _nome_valido(m.group('nome'))

def _val_org(texto, m):
    return _org_valida(m.group('nome'))


# ---------------------------------------------------------------------------
# Padroes de deteccao
# Cada item: (categoria, regex, flags, nome_do_grupo_ou_None, validador_ou_None)
# O validador recebe (texto, match) e retorna bool.
# ---------------------------------------------------------------------------

PADROES = [
    # ------------------------------ DOCUMENTO ------------------------------
    ('DOCUMENTO', r'\b(?:CNH|Carteira Nacional de Habilita[çc][ãa]o)\s*(?:n[ºo°.]?\s*)?\d{9,11}\b', re.I, None, None),
    # RG por gatilho: carteira de identidade n. 134.226 / RG n. 09738019-0
    ('DOCUMENTO', r'\b(?:carteira\s+de\s+identidade|c[eé]dula\s+de\s+identidade|RG|R\.G\.|Registro\s+Geral|Identidade)\s*(?:n[ºo°.:]?\s*)*\d[\d.\-]{2,13}\d?\b', re.I, None, None),
    ('DOCUMENTO', r'(?:expedida?\s+pelo\s+Detran[/\-A-Z]*)\s*0?\d{8,9}\b', re.I, None, None),
    ('DOCUMENTO', r'(?<!\d)\d{11}(?!\d)', 0, None, None),                       # CPF sem formatacao
    ('DOCUMENTO', r'\b(?:PIS|PASEP|NIT|PIS/PASEP)\s*(?:n[o.]?\s*)?\d{3}\.?\d{5}\.?\d{2}-?\d\b', re.I, None, None),
    ('DOCUMENTO', r'\b(?:T[ií]tulo de Eleitor|T[ií]tulo Eleitoral)\s*(?:n[o.]?\s*)?\d{4}\s?\d{4}\s?\d{4}\b', re.I, None, None),
    ('DOCUMENTO', r'\b\d{2}\.?\d{3}\.?\d{3}/\d{4}-?\d{2}\b', 0, None, None),    # CNPJ
    ('DOCUMENTO', r'(?<!\w)[\*\d]{3}\s*\.?\s*[\*\d]{3}\s*\.?\s*[\*\d]{3}\s*-\s*[\*\d]{2}(?!\w)', 0, None, None),
    ('DOCUMENTO', r'\b\d{3}\.?\d{3}\.?\d{3}-?\d{2}\b', 0, None, None),          # CPF formatado
    ('DOCUMENTO', r'\bNIRE\s*(?:n[ºo°.]?\s*)?\d[\d.\-]{4,14}\b', re.I, None, None),
    ('DOCUMENTO', r'\bOAB(?:/[A-Z]{2})?\s*(?:n[oº°.]?\s*)?\d{1,7}(?:/[A-Z]{2})?\b', re.I, None, None),
    ('DOCUMENTO', r'\bOAB\s+[A-Z]{2}\d{5,7}\b', 0, None, None),
    ('DOCUMENTO', r'\b' + _UFS + r'\d{5,7}\b', 0, None, None),                  # registro colado: RJ118454
    ('DOCUMENTO', r'\(\s*\d{8,12}\s*,\s*pg\.?\s*\d+\s*/\s*\d+\s*\)', 0, None, None),
    ('DOCUMENTO', r'\b(?:(?:Justificativa|Parecer|Despacho|Documento|SEI|ID|Id|Protocolo)\s*n?\.?\s*)\(?\d{8,15}\)?\b', 0, None, None),
    ('DOCUMENTO', r'\b(?:Justificativa|Parecer|Despacho|Of[íi]cio|Memorando|Informa[çc][ãa]o|Relat[óo]rio)\s+n?\.?\s*\d{8,15}\b', re.I, None, None),
    ('DOCUMENTO', r'\(\s*\d{8,15}\s*\)', 0, None, None),
    ('DOCUMENTO', r'\b(?:Id|ID|id)\.?\s*\d{8,15}\b', 0, None, None),
    # ------------------------------ PROCESSO -------------------------------
    ('PROCESSO', r'\b\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}\b', 0, None, None),
    ('PROCESSO', r'\b(?:Processo\s+Administrativo\s*)?(?:n[o.]?\s*)?\d{4}\.\d{6}/\d{4}-\d{2}\b', re.I, None, None),
    # ------------------------------ CONTATO --------------------------------
    ('CONTATO', r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', 0, None, None),
    ('CONTATO', r'(?:\(\d{2}\)\s*|\b\d{2}\s+)?(?:9\s*)?\d{4,5}[-\s]\d{4}\b', 0, None, None),
    ('CONTATO', r'https?://[^\s\n,;)>"]+', re.I, None, None),
    ('CONTATO', r'www\.[a-zA-Z0-9\-]+(?:\.[a-zA-Z0-9\-]+)+(?:/[^\s\n,;)>"]*)?', re.I, None, None),
    # ------------------------------ ENDERECO -------------------------------
    ('ENDERECO', r'(?<![\d.])(?:\d{2}[.,]\d{3}-\d{3}|\d{5}-?\d{3})(?!\.?\d)', 0, None, None),   # CEP
    ('ENDERECO', r'\b(?:IP|endere[çc]o\s+IP)\s*(?:\d{1,3}\.){3}\d{1,3}\b', re.I, None, None),
    ('ENDERECO', r'\bmatr[íi]cula(?:\s+do\s+im[óo]vel|\s+imobili[áa]ria)?\s*(?:n[o.]?\s*)?\d{2,10}(?:[./-]\d{1,6})?\b', re.I, None, None),
    # Logradouro completo ate o CEP (ou ate delimitador forte)
    ('ENDERECO', r'\b' + _LOGRADOUROS + r'[ \t]+[^\n]{3,160}?(?=,?[ \t]*[-–—]?[ \t]*CEP\b|,[ \t]*nesta\b|,[ \t]*com[ \t]+sede|,[ \t]*na[ \t]+cidade|;|\n|$)', re.I, None, None),
    # Fallback: logradouro + numero (linhas longas sem CEP)
    ('ENDERECO', r'\b' + _LOGRADOUROS + r'[ \t]+[^\n,;]{2,60}(?:,[ \t]*(?:n[ºo°.]?|s/n)[ \t]*[\w./\- ]{1,15})?', re.I, None, None),
    # Cidade, UF  ou  Cidade/UF (palavras em caixa mista, sem IGNORECASE)
    ('ENDERECO', r'\b' + _PAL_MISTA + r'(?:[ \t]+(?:' + _CONECT_MIN + r'|' + _PAL_MISTA + r')){0,3}[ \t]*[,/][ \t]*' + _UFS + r'\b', 0, None, None),
    # --------------------------- DADOS_BANCARIOS ---------------------------
    ('DADOS_BANCARIOS', r'\b(?:chave\s+)?PIX\s*(?::|-)?\s*(?:CPF|CNPJ|e-mail|email|telefone|celular|aleat[óo]ria)?\s*[A-Za-z0-9._%+\-@]{5,80}\b', re.I, None, None),
    ('DADOS_BANCARIOS', r'\bAg[êe]ncia\s*(?:n[o.]?\s*)?\d{3,6}-?\d?\b', re.I, None, None),
    ('DADOS_BANCARIOS', r'\bConta(?:\s+corrente|\s+poupan[çc]a)?\s*(?:n[o.]?\s*)?\d{4,12}-?\d?\b', re.I, None, None),
    ('DADOS_BANCARIOS', r'\bBanco\s+\d{3}\b', re.I, None, None),
    # ----------------------------- DADO_VEICULO ----------------------------
    ('DADO_VEICULO', r'(?<!: )\b(?:placa\s*)?[A-Z]{3}-\d{4}\b', 0, None, None),
    ('DADO_VEICULO', r'\bplaca\s*[A-Z]{3}\d{4}\b', re.I, None, None),
    ('DADO_VEICULO', r'\b(?:placa\s*)?[A-Z]{3}\d[A-Z]\d{2}\b', 0, None, None),
    # ----------------------------- DADO_SENSIVEL ---------------------------
    ('DADO_SENSIVEL', r'\b(?:Nomeado|Designado|Eleito)\s+(?:atrav[eé]s|por|pelo|nos\s+termos)\s+(?:do\s+)?(?:Decreto|Portaria)[\s\S]{1,80}?(?:\.|\n|\Z)', re.I, None, None),
    # ------------------------------- VALOR ---------------------------------
    # R$ com valor por extenso entre parenteses (evita vazamento pelo extenso)
    ('VALOR', r'R\$[ \t]*\d{1,3}(?:\.\d{3})*(?:,\d{2})?[ \t]*\([^()\n]{1,140}?(?:real|reais|centavos?)\s*\)', re.I, None, None),
    ('VALOR', r'R\$[ \t]*\d{1,3}(?:\.\d{3})*(?:,\d{2})?', 0, None, None),
    # Quantidade com extenso: 135.700 (cento e trinta e cinco mil...) e
    # tambem sem separador de milhar: 700 (setecentas)
    ('VALOR', r'\b\d{1,3}(?:\.\d{3})*(?:,\d{2})?[ \t]*\([^()\n]{1,140}?(?:mil|cem|cent[oa]s?|reais|real)[^()\n]{0,60}?\)', re.I, None, None),
    ('VALOR', r'\b\d{1,3}(?:\.\d{3}){1,4},\d{2}\b', 0, None, None),
    # Numero de milhar isolado (quantidade de quotas em tabelas) — com guarda anti-lei
    ('VALOR', r'\b\d{1,3}(?:\.\d{3}){1,4}\b', 0, None, _num_nao_e_referencia_legal),
    # -------------------------------- DATA ---------------------------------
    # (comente estas linhas para preservar datas nos documentos)
    ('DATA', r'\b\d{1,2}/\d{1,2}/\d{2,4}\b', 0, None, None),
    ('DATA', r'\b\d{1,2}[ºo°]?\s+de\s+' + _MESES + r'\s+de\s+\d{2,4}\b', re.I, None, None),
    ('DATA', r'\b\d{1,2}[ºo°]?\s+de\s+' + _MESES + r'\b', re.I, None, None),
    ('DATA', _MESES + r'\s+de\s+\d{4}\b', re.I, None, None),
    # ---------------------- NOME: organizacoes -----------------------------
    # Categoria interna NOME_ORG: publicada como NOME, mas distingue pessoa de
    # organizacao na propagacao (orgs NAO propagam a primeira palavra isolada).
    # Razao social com sufixo societario (aceita acentos: PRIMAVERA PARTICIPAÇÕES...)
    ('NOME_ORG', r'\b(?P<nome>[A-ZÀ-ÖØ-Þ][\wÀ-ÖØ-öø-ÿ&.\-]*(?:[ \t]+(?:' + _CONECT_MIN + r'|' + _CONECT_CAPS + r'|&|[A-ZÀ-ÖØ-Þ][\wÀ-ÖØ-öø-ÿ&.\-]*)){0,8}[ \t]+' + _SUFIXOS_SOC + r'\b\.?)', 0, 'nome', _val_org),
    # Razao social antes de ", CNPJ" / "inscrita no CNPJ"
    ('NOME_ORG', r'(?P<nome>[A-ZÀ-ÖØ-Þ][\wÀ-ÖØ-öø-ÿ&.\-]*(?:[ \t]+(?:' + _CONECT_MIN + r'|' + _CONECT_CAPS + r'|&|[A-ZÀ-ÖØ-Þ][\wÀ-ÖØ-öø-ÿ&.\-]*)){1,8})(?=[ \t]*,?[ \t]*(?:inscrit[oa][ \t]+no[ \t]+)?(?:CNPJ|CNPJ/MF)\b)', 0, 'nome', _val_org),
    # Sociedade de advogados: SAMPAIO E CAMPOS ADVOGADOS / Silva Advocacia
    ('NOME_ORG', r'\b(?P<nome>[A-ZÀ-ÖØ-Þ][\wÀ-ÖØ-öø-ÿ&.\-]*(?:[ \t]+(?:' + _CONECT_MIN + r'|' + _CONECT_CAPS + r'|&|[A-ZÀ-ÖØ-Þ][\wÀ-ÖØ-öø-ÿ&.\-]*)){0,6}[ \t]+' + _SUFIXOS_ADV + r'\b)', 0, 'nome', _val_org),
    # Fundacao/Instituto/Autarquia + nome
    ('NOME_ORG', r'\b(?P<nome>(?:Funda[çc][ãa]o|Instituto|Autarquia)(?:[ \t]+[A-Za-zÀ-ÖØ-öø-ÿ]+){2,15}(?:[ \t]*-[ \t]*[A-Z0-9-]{3,30})?)', 0, 'nome', _val_org),
    # ---------------------- NOME: pessoas fisicas --------------------------
    # Caixa alta + qualificacao: ANA MARIA DE ANDRADE CAMPOS, brasileira, ...
    ('NOME', r'\b(?P<nome>' + _NOME_CAPS + r')(?=[ \t]*,[ \t]*' + _QUALIFICADORES + r')', 0, 'nome', _val_nome),
    # Caixa mista + qualificacao: Ana Maria de Andrade Campos, brasileira, ...
    ('NOME', r'\b(?P<nome>' + _NOME_MISTO + r')(?=[ \t]*,[ \t]*' + _QUALIFICADORES + r')', 0, 'nome', _val_nome),
    # Apos pronome de tratamento: Sr./Sra./Dr./Dra. + Nome
    ('NOME', r'\b(?:Sr\.?|Sra\.?|Senhor|Senhora|Dr\.?|Dra\.?)[ \t]+(?P<nome>' + _NOME_MISTO + r'|' + _NOME_CAPS + r')', 0, 'nome', _val_nome),
    # Apos papel processual: autor/reu/requerente + Nome
    ('NOME', r'\b(?i:autora?|r[eé]u?|requerente|requerid[oa])[ \t]+(?P<nome>' + _NOME_MISTO + r'|' + _NOME_CAPS + r')', 0, 'nome', _val_nome),
    # representada por / neste ato representado por + NOME
    ('NOME', r'\b(?i:representad[oa][ \t]+(?:neste[ \t]+ato[ \t]+)?por)[ \t]+(?:(?i:(?:seu|sua)[ \t]+\w+[ \t]+))?(?:Sr\.?|Sra\.?|Dr\.?|Dra\.?)?[ \t]*(?P<nome>' + _NOME_CAPS + r'|' + _NOME_MISTO + r')', 0, 'nome', _val_nome),
    # Linha isolada em caixa alta (assinaturas): ANA MARIA DE ANDRADE CAMPOS
    ('NOME', r'(?m)^[ \t]*(?P<nome>' + _NOME_CAPS + r')[ \t]*\r?$', 0, 'nome', _val_nome),
    # Linha isolada em caixa mista: Beatriz Andrade Campos
    ('NOME', r'(?m)^[ \t]*(?P<nome>' + _NOME_MISTO + r')[ \t]*\r?$', 0, 'nome', _val_nome),
    # Nome antes de virgula + OAB (caixa mista): Andrade Campos, OAB/RJ n.
    ('NOME', r'(?P<nome>' + _NOME_MISTO + r')(?=[ \t]*,[ \t]*OAB\b)', 0, 'nome', _val_nome),
    # Gatilhos de cabecalho: RELATOR: NOME / AUTOR: NOME / Paciente: Nome
    ('NOME', r'(?i:RELATORA?|JUIZA?|MAGISTRADO|AUTORA?|R[EÉ]U?|Recolhid[oa][ \t]+por|Emitid[oa][ \t]+para|Titular|Segurad[oa]|Paciente|Benefici[áa]ri[oa]|Dependente)\s*:\s*(?P<nome>' + _NOME_CAPS + r'|' + _NOME_MISTO + r')', 0, 'nome', _val_nome),
    # Assinatura eletronica: CN=NOME SOBRENOME
    ('NOME', r'(?i:CN)=(?P<nome>' + _NOME_CAPS + r')', 0, 'nome', _val_nome),
    # Documento assinado eletronicamente por NOME,
    ('NOME', r'(?i:assinado\s+eletronicamente\s+por)\s+(?P<nome>' + _NOME_MISTO + r'|' + _NOME_CAPS + r')(?=\s*,)', 0, 'nome', _val_nome),
    # Eu, NOME SOBRENOME,
    ('NOME', r'\b(?i:Eu),\s+(?P<nome>' + _NOME_MISTO + r'|' + _NOME_CAPS + r')(?=\s*,)', 0, 'nome', _val_nome),
    # substabeleco ..., NOME,
    ('NOME', r'(?i:substabele[çc]o[^,\n]{0,60},\s+(?:com\s+reserva\s+de\s+iguais,?\s+)?)(?P<nome>' + _NOME_MISTO + r'|' + _NOME_CAPS + r')(?=\s*,)', 0, 'nome', _val_nome),
    # OUTORGANTE: NOME,
    ('NOME', r'(?i:OUTORGANTE):\s*(?P<nome>' + _NOME_MISTO + r'|' + _NOME_CAPS + r')(?=\s*,)', 0, 'nome', _val_nome),
    # Blocos cadastrais eproc: NOME SOBRENOME   RJ123456 / NOME   (CPF...)
    ('NOME', r'(?m)^[ \t]*(?P<nome>' + _NOME_CAPS + r')(?=[ \t]{2,}(?:' + _UFS + r'\d{5,7}|\(|\d))', 0, 'nome', _val_nome),
    # RJ123456 - NOME SOBRENOME - ADVOGADO/MAGISTRADO
    ('NOME', r'(?m)^(?:[A-Z]{2}\d{5,7}|[A-Z]{5,20})\s*-\s*(?P<nome>' + _NOME_CAPS + r')\s*-\s*(?:ADVOGADO|MAGISTRADO|JUIZ|PROMOTOR|DEFENSOR)', 0, 'nome', _val_nome),
    # ADVOGADO(A): NOME SOBRENOME (OAB ...)
    ('NOME', r'(?i:ADVOGADO|ADVOGADA)\s*\(?(?i:A)?\)?\s*:\s*(?P<nome>' + _NOME_CAPS + r')(?=\s*\()', 0, 'nome', _val_nome),
    # Dois nomes em caixa alta separados por virgula em linha isolada
    ('NOME', r'(?m)^[ \t]*(?P<nome>' + _NOME_CAPS + r')(?=,[ \t]*' + _NOME_CAPS + r'[ \t]*\r?$)', 0, 'nome', _val_nome),
    ('NOME', r'(?m)^[ \t]*' + _NOME_CAPS + r',[ \t]*(?P<nome>' + _NOME_CAPS + r')[ \t]*\r?$', 0, 'nome', _val_nome),
    # NOME SOBRENOME assinou como contratante
    ('NOME', r'(?P<nome>' + _NOME_CAPS + r')(?=[ \t]+assinou[ \t]+como)', 0, 'nome', _val_nome),
    # NOME SOBRENOME, devidamente qualificado nos autos
    ('NOME', r'(?P<nome>' + _NOME_CAPS + r')(?=\s*,\s*devidamente[ \t]+qualificad)', 0, 'nome', _val_nome),
]

PADROES_COMPILADOS = [
    (cat, re.compile(rx, flags), grupo, val)
    for (cat, rx, flags, grupo, val) in PADROES
]


# ---------------------------------------------------------------------------
# Registro de entidades (numeracao consistente + mapa de reversao)
# ---------------------------------------------------------------------------

class Registro:
    """
    Gera tokens numerados por categoria: [NOME_01], [DOCUMENTO_01]...
    O mesmo dado (normalizado) recebe SEMPRE o mesmo token no documento.
    Armazena token -> dado_original para reversao (cofre de chaves).

    modelos: dict opcional categoria -> modelo de token com placeholders
             {cat} e {n} (ex.: '[{cat}_{n:02d}]', 'PARTE {n}').
    """
    MODELO_PADRAO = '[{cat}_{n:02d}]'

    def __init__(self, modelos=None):
        self._seq = {}
        self._mapa = {}       # token -> dado_original (primeira ocorrencia)
        self._por_chave = {}  # (categoria, dado_normalizado) -> token
        self._modelos = dict(modelos or {})
        self.substituicoes = {}

    def token(self, categoria, dado):
        categoria = _cat_publica(categoria)
        chave = (categoria, _norm_entidade(dado))
        tok = self._por_chave.get(chave)
        if tok is None:
            n = self._seq.get(categoria, 0) + 1
            self._seq[categoria] = n
            modelo = self._modelos.get(categoria) or self.MODELO_PADRAO
            try:
                tok = modelo.format(cat=categoria, n=n)
            except (KeyError, ValueError, IndexError):
                tok = self.MODELO_PADRAO.format(cat=categoria, n=n)
            self._por_chave[chave] = tok
            self._mapa[tok] = dado
        self.substituicoes[categoria] = self.substituicoes.get(categoria, 0) + 1
        return tok

    @property
    def contagens(self):
        return dict(self.substituicoes)

    @property
    def mapa(self):
        return dict(self._mapa)


# ---------------------------------------------------------------------------
# Motor de spans
# ---------------------------------------------------------------------------

def _coletar_spans(texto):
    """Detecta todos os padroes sobre o texto original. Retorna [(ini, fim, cat, dado)]."""
    spans = []
    for categoria, rx, grupo, validador in PADROES_COMPILADOS:
        for m in rx.finditer(texto):
            if validador and not validador(texto, m):
                continue
            if grupo:
                ini, fim = m.span(grupo)
                dado = m.group(grupo)
            else:
                ini, fim = m.span()
                dado = m.group(0)
            if ini < fim:
                spans.append((ini, fim, categoria, dado))
    return spans


def _resolver_sobreposicoes(spans):
    """Ordena e remove sobreposicoes: o span mais longo que comeca antes vence."""
    spans = sorted(spans, key=lambda s: (s[0], -(s[1] - s[0])))
    resultado, fim_anterior = [], -1
    for s in spans:
        if s[0] >= fim_anterior:
            resultado.append(s)
            fim_anterior = s[1]
    return resultado


_CONECTIVOS = {'de', 'da', 'do', 'dos', 'das', 'e'}


def _variantes_prefixo(nome, tam_min_primeira):
    """
    Variantes de referencia a um nome ja detectado, para propagar mencoes curtas:
      'ANA MARIA DE ANDRADE CAMPOS'                -> 'ANA MARIA' e 'ANA'
      'PRIMAVERA PARTICIPACOES E INVESTIMENTOS'    -> 'PRIMAVERA PARTICIPACOES' e 'PRIMAVERA'
    Guardas: primeira palavra com tamanho minimo e fora das stopwords/genericos.
    """
    palavras = nome.split()
    if len([p for p in palavras if p.lower() not in _CONECTIVOS]) < 2:
        return []
    primeira = palavras[0]
    chave = _sem_acento(primeira).upper()
    if (len(primeira) < tam_min_primeira
            or chave in STOPWORDS_NOME or chave in _GENERICOS_ORG):
        return []
    variantes = []
    if palavras[1].lower() not in _CONECTIVOS:
        variantes.append(primeira + ' ' + palavras[1])   # Ana Maria / Primavera Sabores
    variantes.append(primeira)                            # Ana / Primavera
    return variantes


_SUFIXOS_GERACAO = {'NETO', 'NETTO', 'FILHO', 'FILHA', 'JUNIOR', 'SOBRINHO'}


def _variantes_sufixo(nome, tam_min=4):
    """
    Variantes pelo FINAL do nome (sobrenomes), para propagar referencias como
    'o Dr. Sampaio' ou 'a familia Silva Prado':
      'ANA MARIA DE ANDRADE CAMPOS' -> 'ANDRADE CAMPOS' e 'CAMPOS'
    Sufixo de geracao (Neto, Filho, Junior) nunca vira variante isolada:
      'MANOEL M. DA SILVA PRADO NETO' -> 'SILVA PRADO NETO' e 'PRADO NETO'.
    Guardas: tamanho minimo e fora das stopwords; iniciais abreviadas ignoradas.
    """
    fortes = [p for p in nome.split()
              if p.lower() not in _CONECTIVOS and not re.fullmatch(_INICIAL, p)]
    if len(fortes) < 2:
        return []

    def ok(p):
        chave = _sem_acento(p).upper()
        return (len(p) >= tam_min and chave not in STOPWORDS_NOME
                and chave not in _GENERICOS_ORG)

    variantes = []
    ultimo = fortes[-1]
    if _sem_acento(ultimo).upper() in _SUFIXOS_GERACAO:
        if len(fortes) >= 3 and ok(fortes[-2]):
            variantes.append(fortes[-2] + ' ' + ultimo)          # Prado Neto
            if len(fortes) >= 4 and ok(fortes[-3]):
                variantes.append(' '.join(fortes[-3:]))          # Silva Prado Neto
    elif ok(ultimo):
        variantes.append(ultimo)                                  # Campos
        if len(fortes) >= 3 and ok(fortes[-2]):
            variantes.append(fortes[-2] + ' ' + ultimo)          # Andrade Campos
    return variantes


def _entidades_propagaveis(listas_de_spans):
    """
    Extrai nomes ja detectados para propagar a TODAS as ocorrencias no documento.
      - pessoas: nome completo + variantes de prenome ('Ana Maria', 'Ana');
      - organizacoes: razao social + variante sem sufixo societario
        (PRIMAVERA SABORES LTDA -> 'PRIMAVERA SABORES') + nome de fantasia
        ('PRIMAVERA PARTICIPACOES' e 'PRIMAVERA').
    Variantes de palavra unica so casam ocorrencias CAPITALIZADAS (ver
    _propagar_entidades), evitando atingir palavras comuns do texto.
    Retorna lista de (valor, categoria_interna), mais longos primeiro.
    """
    entidades, vistos = [], set()

    def adicionar(valor, cat, minimo=7, min_palavras=2):
        valor = ' '.join(valor.split()).strip(' .,;:')
        if len(valor) < minimo or len(valor.split()) < min_palavras:
            return
        chave = _norm_entidade(valor)
        if chave not in vistos:
            vistos.add(chave)
            entidades.append((valor, cat))

    for spans in listas_de_spans:
        for (_, _, cat, dado) in spans:
            dado_limpo = ' '.join(dado.split())
            if cat == 'NOME':          # pessoa fisica
                adicionar(dado, 'NOME')
                for v in _variantes_prefixo(dado_limpo, tam_min_primeira=3):
                    adicionar(v, 'NOME', minimo=3, min_palavras=1)
                for v in _variantes_sufixo(dado_limpo):
                    adicionar(v, 'NOME', minimo=4, min_palavras=1)
            elif cat == 'NOME_ORG':    # organizacao
                adicionar(dado, 'NOME_ORG')
                base = dado_limpo
                m = re.match(r'(.+?)[ \t]+(?:' + _SUFIXOS_SOC + r'|' + _SUFIXOS_ADV + r')\.?\s*$', dado_limpo)
                if m:
                    base = m.group(1)
                    adicionar(base, 'NOME_ORG')
                for v in _variantes_prefixo(base, tam_min_primeira=4):
                    adicionar(v, 'NOME_ORG', minimo=4, min_palavras=1)

    entidades.sort(key=lambda ec: len(ec[0]), reverse=True)   # mais longas primeiro
    return entidades


def _propagar_entidades(texto, spans, entidades):
    """
    Adiciona spans para toda ocorrencia literal (case-insensitive) das entidades.
    Variantes de PALAVRA UNICA ('Ana', 'Primavera') so casam quando a ocorrencia
    esta capitalizada — referencias a pessoa/empresa em texto juridico sempre
    estao; assim o substantivo comum ('durante a primavera') fica intacto.
    """
    if not entidades:
        return spans
    ocupados = [(s[0], s[1]) for s in spans]

    def livre(ini, fim):
        return all(fim <= a or ini >= b for (a, b) in ocupados)

    extras = []
    for ent, cat in entidades:
        palavra_unica = ' ' not in ent
        rx = re.compile(
            r'(?<![\wÀ-ÖØ-öø-ÿ])' + r'\s+'.join(re.escape(p) for p in ent.split()) + r'(?![\wÀ-ÖØ-öø-ÿ])',
            re.I,
        )
        for m in rx.finditer(texto):
            if palavra_unica and not m.group(0)[:1].isupper():
                continue
            if livre(m.start(), m.end()):
                extras.append((m.start(), m.end(), cat, ent))
                ocupados.append((m.start(), m.end()))
    return _resolver_sobreposicoes(spans + extras)


def _aplicar_spans_texto(texto, spans_tokenizados):
    """Aplica [(ini, fim, token)] (ordenados, sem sobreposicao) em uma passada."""
    partes, pos = [], 0
    for ini, fim, token in spans_tokenizados:
        partes.append(texto[pos:ini])
        partes.append(token)
        pos = fim
    partes.append(texto[pos:])
    return ''.join(partes)


def anonimizar_texto(texto):
    """
    Anonimiza texto plano (txt/pdf). Retorna (texto_anon, contagens, mapa_reversao).
    """
    reg = Registro()
    spans = _resolver_sobreposicoes(_coletar_spans(texto))
    entidades = _entidades_propagaveis([spans])
    spans = _propagar_entidades(texto, spans, entidades)
    spans_tok = [(i, f, reg.token(c, d)) for (i, f, c, d) in spans]
    return _aplicar_spans_texto(texto, spans_tok), reg.contagens, reg.mapa


# ---------------------------------------------------------------------------
# DOCX: iteracao completa (paragrafos + tabelas + cabecalhos/rodapes)
# ---------------------------------------------------------------------------

def _iter_paragrafos_docx(doc):
    """Percorre todos os paragrafos do documento, inclusive tabelas aninhadas,
    cabecalhos e rodapes. Deduplica celulas mescladas pelo elemento XML da
    celula (mantido no set para que a identidade do proxy lxml fique estavel;
    id() puro nao serve, pois o lxml recicla proxies e os ids se repetem)."""
    celulas_vistas = set()

    def de_tabela(tabela):
        for linha in tabela.rows:
            for celula in linha.cells:
                tc = celula._tc
                if tc in celulas_vistas:
                    continue
                celulas_vistas.add(tc)
                yield from celula.paragraphs
                for t in celula.tables:
                    yield from de_tabela(t)

    yield from doc.paragraphs
    for t in doc.tables:
        yield from de_tabela(t)
    for secao in doc.sections:
        for parte in (secao.header, secao.footer):
            yield from parte.paragraphs
            for t in parte.tables:
                yield from de_tabela(t)


def _aplicar_spans_runs(para, spans_tokenizados):
    """
    Aplica [(ini, fim, token)] no paragrafo, run a run, por OFFSET exato.
    Preserva a formatacao: o token assume o estilo do run onde o span comeca.
    Processa de tras para frente para nao invalidar offsets.
    """
    if not spans_tokenizados:
        return
    runs = para.runs
    textos = [r.text for r in runs]

    if not runs or ''.join(textos) != para.text:
        # Estrutura atipica (hyperlinks etc.): reescreve o paragrafo inteiro.
        novo = _aplicar_spans_texto(para.text, spans_tokenizados)
        if runs:
            runs[0].text = novo
            for r in runs[1:]:
                r.text = ''
        else:
            para.text = novo
        return

    offsets, pos = [], 0
    for t in textos:
        offsets.append(pos)
        pos += len(t)

    for ini, fim, token in reversed(spans_tokenizados):
        i = bisect.bisect_right(offsets, ini) - 1
        j = bisect.bisect_right(offsets, max(ini, fim - 1)) - 1
        rel_i = ini - offsets[i]
        rel_j = fim - offsets[j]
        if i == j:
            textos[i] = textos[i][:rel_i] + token + textos[i][rel_j:]
        else:
            resto = textos[j][rel_j:]
            textos[i] = textos[i][:rel_i] + token
            for k in range(i + 1, j):
                textos[k] = ''
            textos[j] = resto

    for r, t in zip(runs, textos):
        if r.text != t:
            r.text = t


def _spans_docx(doc):
    """Fase de deteccao completa de um .docx: spans por paragrafo, ja com
    sobreposicoes resolvidas e entidades propagadas."""
    paras = list(_iter_paragrafos_docx(doc))
    spans_por_par = [_resolver_sobreposicoes(_coletar_spans(p.text)) for p in paras]
    entidades = _entidades_propagaveis(spans_por_par)
    spans_por_par = [_propagar_entidades(p.text, s, entidades)
                     for p, s in zip(paras, spans_por_par)]
    return paras, spans_por_par


def _detectar(caminho_entrada, ext):
    """
    Deteccao deterministica para qualquer formato suportado.
    Retorna (tipo, contexto, listas_de_spans):
      tipo 'docx'  -> contexto = (doc, paras)
      tipo 'texto' -> contexto = texto completo (txt ou pdf com _SEP_PAGINA)
    """
    if ext == '.docx':
        if not DOCX_DISPONIVEL:
            raise RuntimeError('python-docx nao instalado. Execute: pip install python-docx')
        doc = _DocxDocument(caminho_entrada)
        paras, spans_por_par = _spans_docx(doc)
        return 'docx', (doc, paras), spans_por_par
    if ext == '.txt':
        texto = _ler_txt(caminho_entrada)
    elif ext == '.pdf':
        if not PDF_DISPONIVEL:
            raise RuntimeError('PyMuPDF nao instalado. Execute: pip install pymupdf')
        texto = _SEP_PAGINA.join(_ler_pdf(caminho_entrada))
    else:
        raise ValueError("Formato nao suportado: '%s'. Use .txt, .docx ou .pdf." % ext)
    spans = _resolver_sobreposicoes(_coletar_spans(texto))
    entidades = _entidades_propagaveis([spans])
    spans = _propagar_entidades(texto, spans, entidades)
    return 'texto', texto, [spans]


# ---------------------------------------------------------------------------
# API de duas fases (usada pelo app local e pela CLI)
# ---------------------------------------------------------------------------

def analisar_arquivo(caminho_entrada, modelos=None):
    """
    Fase 1 (revisao): detecta as entidades SEM gravar nada.
    Retorna lista de entidades na ordem de aparicao no documento:
      [{'token', 'categoria', 'dado', 'ocorrencias'}]
    'token' e a substituicao default, gerada pelos modelos por categoria;
    a deteccao e deterministica: aplicar_anonimizacao() com os mesmos
    modelos reproduz exatamente os mesmos tokens.
    """
    ext = os.path.splitext(caminho_entrada)[1].lower()
    _, _, listas = _detectar(caminho_entrada, ext)
    reg = Registro(modelos)
    entidades, indice = [], {}
    for spans in listas:
        for (_, _, cat, dado) in spans:
            tok = reg.token(cat, dado)
            ent = indice.get(tok)
            if ent is None:
                ent = {'token': tok, 'categoria': _cat_publica(cat),
                       'dado': reg.mapa[tok], 'ocorrencias': 0}
                indice[tok] = ent
                entidades.append(ent)
            ent['ocorrencias'] += 1
    return entidades


def aplicar_anonimizacao(caminho_entrada, caminho_saida, caminho_mapa,
                         substituicoes=None, excluidos=None, modelos=None,
                         compliance=None, permitir_ambiguas=False):
    """
    Fase 2: aplica a anonimizacao e grava saida + mapa de reversao (cofre).
      substituicoes: dict token_default -> texto customizado de substituicao
      excluidos: tokens default que NAO devem ser substituidos (falsos positivos)
      modelos: modelos de token por categoria (os MESMOS usados na analise)
      permitir_ambiguas: permite o MESMO texto de substituicao para dados
        diferentes (ex.: tudo 'XXX'). Nesse caso os textos repetidos sao
        aplicados normalmente, mas ficam FORA do mapa de reversao — esses
        trechos nao poderao ser restaurados (regressao parcial). Com False
        (padrao), a ambiguidade gera ValueError.
    O mapa de reversao usa como chave o texto efetivamente gravado no documento
    (default ou customizado), de modo que a restauracao funciona nos dois casos.
    Para entrada .pdf ou .txt, o formato de saida e definido pela extensao de
    caminho_saida (.pdf, .txt ou .docx). Entrada .docx sempre gera .docx.
    Retorna o relatorio (dict) no mesmo formato da CLI.
    """
    substituicoes = dict(substituicoes or {})
    excluidos = set(excluidos or ())
    ext = os.path.splitext(caminho_entrada)[1].lower()
    hash_original = _calcular_hash(caminho_entrada)

    tipo, contexto, listas = _detectar(caminho_entrada, ext)
    if tipo == 'docx':
        texto_original = '\n'.join(p.text for p in contexto[1])
    else:
        texto_original = contexto

    reg = Registro(modelos)
    contagens, mapa_final, usados = {}, {}, {}
    ambiguas = set()

    def traduzir(spans):
        out = []
        for (i, f, cat, dado) in spans:
            tok = reg.token(cat, dado)
            if tok in excluidos:
                continue
            texto_sub = substituicoes.get(tok, tok)
            dado_canonico = reg.mapa[tok]
            anterior = usados.get(texto_sub)
            if anterior is not None and _norm_entidade(anterior) != _norm_entidade(dado_canonico):
                if not permitir_ambiguas:
                    raise ValueError(
                        'Substituicao ambigua: o texto %r foi atribuido a dados diferentes '
                        '(%r e %r). Use textos distintos para permitir a reversao, ou '
                        'autorize explicitamente a regressao parcial.'
                        % (texto_sub, anterior, dado_canonico))
                ambiguas.add(texto_sub)
            usados.setdefault(texto_sub, dado_canonico)
            mapa_final[texto_sub] = dado_canonico
            cat_pub = _cat_publica(cat)
            contagens[cat_pub] = contagens.get(cat_pub, 0) + 1
            out.append((i, f, texto_sub))
        return out

    if tipo == 'docx':
        doc, paras = contexto
        for para, spans in zip(paras, listas):
            _aplicar_spans_runs(para, traduzir(spans))
        doc.save(caminho_saida)
    else:
        # O formato de saida e definido pela EXTENSAO de caminho_saida:
        # um .pdf de entrada pode gerar .pdf, .txt ou .docx anonimizado.
        texto_anon = _aplicar_spans_texto(contexto, traduzir(listas[0]))
        _gravar_texto_plano(caminho_saida, texto_anon)

    # Textos ambiguos (mesmo texto para dados diferentes) sao IRREVERSIVEIS:
    # saem do mapa de reversao — a restauracao os ignora e eles permanecem
    # anonimizados no arquivo restaurado.
    for texto_sub in ambiguas:
        mapa_final.pop(texto_sub, None)

    # Avisos: texto customizado que ja existia no documento original
    # compromete a reversao (a restauracao substituiria texto legitimo).
    avisos = []
    if ambiguas:
        avisos.append(
            'Substituicoes IRREVERSIVEIS (mesmo texto para dados diferentes, '
            'fora do mapa de reversao): %s. Esses trechos nao serao restaurados '
            'na regressao.' % ', '.join(repr(t) for t in sorted(ambiguas)))
    for texto_sub in substituicoes.values():
        if texto_sub and texto_sub not in ambiguas and texto_sub in texto_original:
            avisos.append(
                'O texto de substituicao %r ja ocorre no documento original; '
                'a restauracao pode substituir texto legitimo.' % texto_sub)

    # Salva o mapa de reversao (cofre de chaves)
    with open(caminho_mapa, 'w', encoding='utf-8') as f:
        json.dump({
            'arquivo_original': os.path.basename(caminho_entrada),
            'hash_sha256_original': hash_original,
            'gerado_em': datetime.now().isoformat(timespec='seconds'),
            'aviso': (
                'CONFIDENCIAL — Este arquivo contem o mapeamento de dados pessoais anonimizados. '
                'Mantenha-o em local seguro e separado dos arquivos anonimizados. '
                'Nao compartilhe junto com os arquivos enviados a redes publicas ou IA.'
            ),
            'mapa_reversao': mapa_final,
        }, f, ensure_ascii=False, indent=2)

    c = compliance or {}
    try:
        operador = getpass.getuser()
    except Exception:
        operador = 'desconhecido'

    return {
        'arquivo_entrada': os.path.basename(caminho_entrada),
        'arquivo_saida': os.path.basename(caminho_saida),
        'arquivo_mapa_reversao': os.path.basename(caminho_mapa),
        'hash_sha256_original': hash_original,
        'processado_em': datetime.now().isoformat(timespec='seconds'),
        'operador': operador,
        'finalidade': c.get('finalidade', 'nao informado'),
        'destinatario': c.get('destinatario', 'nao informado'),
        'base_legal': c.get('base_legal', 'nao informado'),
        'categoria_titular': c.get('categoria_titular', 'nao informado'),
        'ocorrencias_por_tipo': contagens,
        'entidades_unicas': len(mapa_final),
        'total_substituicoes': sum(contagens.values()),
        'substituicoes_customizadas': len(substituicoes),
        'entidades_excluidas': len(excluidos),
        'substituicoes_irreversiveis': sorted(ambiguas),
        'avisos': avisos,
    }


# ---------------------------------------------------------------------------
# Reanominizacao (restauracao)
# ---------------------------------------------------------------------------

_TOKEN_RX = re.compile(r'\[[A-Z_]+_\d{2,}\]')


def _resolver_mapa(mapa_reversao, max_iter=10):
    """Resolve tokens aninhados no mapa de reversao (heranca de mapas v3)."""
    mapa = dict(mapa_reversao)
    for _ in range(max_iter):
        alterou = False
        for token, valor in mapa.items():
            novo = valor
            for sub in _TOKEN_RX.findall(valor):
                if sub in mapa and sub != token:
                    novo = novo.replace(sub, mapa[sub])
            if novo != valor:
                mapa[token] = novo
                alterou = True
        if not alterou:
            break
    return mapa


def reanominizar_texto(texto_anonimizado, mapa_reversao):
    """Restaura os dados originais em um texto anonimizado usando o mapa."""
    mapa = _resolver_mapa(mapa_reversao)
    t = texto_anonimizado
    for token, dado in sorted(mapa.items(), key=lambda x: -len(x[0])):
        t = t.replace(token, dado)
    return t


def _spans_de_tokens(texto, mapa):
    """
    Localiza ocorrencias LITERAIS das chaves do mapa de reversao no texto
    (chaves maiores primeiro, sem sobreposicao). Funciona tanto com tokens
    default ([NOME_01]) quanto com textos de substituicao customizados.
    """
    spans, ocupados = [], []
    for token in sorted(mapa, key=len, reverse=True):
        if not token:
            continue
        ini = texto.find(token)
        while ini >= 0:
            fim = ini + len(token)
            if all(fim <= a or ini >= b for (a, b) in ocupados):
                spans.append((ini, fim, mapa[token]))
                ocupados.append((ini, fim))
            ini = texto.find(token, fim)
    return sorted(spans)


def _restaurar_docx(caminho_anonimizado, caminho_saida, mapa_reversao):
    """Restaura um .docx anonimizado substituindo tokens (default ou
    customizados) pelos dados originais, preservando formatacao.
    Retorna o numero de trechos efetivamente revertidos."""
    mapa = _resolver_mapa(mapa_reversao)
    doc = _DocxDocument(caminho_anonimizado)
    revertidos = 0
    for para in _iter_paragrafos_docx(doc):
        spans = _spans_de_tokens(para.text, mapa)
        revertidos += len(spans)
        _aplicar_spans_runs(para, spans)
    doc.save(caminho_saida)
    return revertidos


# ---------------------------------------------------------------------------
# LGPD: hash, config e log consolidado
# ---------------------------------------------------------------------------

def _calcular_hash(caminho):
    h = hashlib.sha256()
    with open(caminho, 'rb') as f:
        for bloco in iter(lambda: f.read(65536), b''):
            h.update(bloco)
    return h.hexdigest()


_CONFIG_PADRAO = {
    "_instrucoes": (
        "Preencha os campos abaixo para conformidade com a LGPD "
        "(art. 37 e art. 6, VIII). Salve o arquivo e execute o "
        "anonimizador normalmente."
    ),
    "finalidade": "PREENCHER: motivo do tratamento (ex: uso em IA generativa, analise interna, compartilhamento com parceiro X)",
    "destinatario": "PREENCHER: destino do arquivo anonimizado (ex: uso interno, OpenAI API, cliente Y)",
    "base_legal": "PREENCHER: hipotese LGPD (ex: Art. 7, V - execucao de contrato; Art. 7, II - consentimento; Art. 7, IX - interesse legitimo)",
    "categoria_titular": "PREENCHER: categoria dos titulares (ex: partes processuais, clientes, funcionarios, terceiros)",
}


CAMPOS_LGPD = ('finalidade', 'destinatario', 'base_legal', 'categoria_titular')


def ler_config(pasta_base):
    """
    Le (ou cria) o config.json SEM abortar — para uso do app local.
    Retorna (config_dict, lista_de_campos_incompletos).
    """
    caminho = os.path.join(pasta_base, 'config.json')
    if not os.path.exists(caminho):
        with open(caminho, 'w', encoding='utf-8') as f:
            json.dump(_CONFIG_PADRAO, f, ensure_ascii=False, indent=2)
        cfg = dict(_CONFIG_PADRAO)
    else:
        with open(caminho, encoding='utf-8') as f:
            cfg = json.load(f)
    campos = {c: cfg.get(c, '') for c in CAMPOS_LGPD}
    incompletos = [c for c, v in campos.items()
                   if not v or v.startswith('PREENCHER')]
    return campos, incompletos


def salvar_config(pasta_base, campos):
    """Grava os campos LGPD no config.json preservando as instrucoes."""
    caminho = os.path.join(pasta_base, 'config.json')
    cfg = {'_instrucoes': _CONFIG_PADRAO['_instrucoes']}
    cfg.update({c: campos.get(c, '') for c in CAMPOS_LGPD})
    with open(caminho, 'w', encoding='utf-8') as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


def _carregar_config(pasta_base):
    """Carrega config.json com fail-fast (CLI) se campos nao preenchidos."""
    campos, incompletos = ler_config(pasta_base)
    caminho = os.path.join(pasta_base, 'config.json')
    if incompletos:
        print('[ERRO LGPD] config.json com campos nao preenchidos: %s' % ', '.join(incompletos))
        print('            Preencha o arquivo antes de executar o anonimizador.')
        print('            Caminho: %s\n' % caminho)
        print('            O processamento foi INTERROMPIDO para proteger a integridade do log de compliance.')
        sys.exit(1)
    return campos


# ---------------------------------------------------------------------------
# Normalizacao de caracteres Unicode
# ---------------------------------------------------------------------------

_UNICODE_MAP = {
    '“': '"', '”': '"',
    '‘': "'", '’': "'",
    '–': '-',
    '—': '--',
    '…': '...',
    ' ': ' ',
    '«': '"', '»': '"',
    # Ligaduras tipograficas geradas pelo PyMuPDF ao extrair PDFs
    'ﬀ': 'ff', 'ﬁ': 'fi', 'ﬂ': 'fl',
    'ﬃ': 'ffi', 'ﬄ': 'ffl',
    '·': 'tp',   # ponto medio usado como ligadura 'tp' (h.ps -> https)
    'Ʃ': 'tt',   # ESH usado como ligadura 'tt' pelo PyMuPDF
    'ﬆ': 'st', 'ﬅ': 'st',
}

def _normalizar_unicode(texto):
    for orig, sub in _UNICODE_MAP.items():
        texto = texto.replace(orig, sub)
    return texto

# Separador de pagina para PDFs (rodeado de \n para os padroes de linha isolada)
_SEP_PAGINA = '\n\x01\n'

# ---------------------------------------------------------------------------
# I/O de arquivos
# ---------------------------------------------------------------------------

def _ler_txt(caminho):
    with open(caminho, encoding='utf-8') as f:
        return f.read()

def _gravar_txt(caminho, texto):
    with open(caminho, 'w', encoding='utf-8') as f:
        f.write(texto)

def _ler_pdf(caminho):
    doc = _fitz.open(caminho)
    paginas = [_normalizar_unicode(page.get_text('text')) for page in doc]
    doc.close()
    if not any(p.strip() for p in paginas):
        raise ValueError(
            'PDF sem texto extraivel. Provavelmente e um PDF escaneado (imagem). '
            'Esta versao suporta apenas PDFs digitais (texto selecionavel). '
            'Para PDFs escaneados, utilize uma ferramenta de OCR antes de processar.'
        )
    return paginas

def _gravar_docx_texto(caminho_saida, paginas_texto):
    """Gera um .docx simples a partir de texto plano (uma pagina por bloco,
    separadas por quebra de pagina). Usado para saida DOCX de origem .pdf/.txt."""
    if not DOCX_DISPONIVEL:
        raise RuntimeError('python-docx nao instalado. Execute: pip install python-docx')
    doc = _DocxDocument()
    for i, texto_pagina in enumerate(paginas_texto):
        if i:
            doc.add_page_break()
        for linha in texto_pagina.split('\n'):
            doc.add_paragraph(linha)
    doc.save(caminho_saida)


def _gravar_texto_plano(caminho_saida, texto_anon):
    """Grava texto plano anonimizado (origem .txt ou .pdf) no formato indicado
    pela EXTENSAO de caminho_saida: .txt, .docx ou .pdf."""
    ext_saida = os.path.splitext(caminho_saida)[1].lower()
    paginas = texto_anon.split(_SEP_PAGINA)
    if ext_saida == '.txt':
        _gravar_txt(caminho_saida, '\n\n'.join(paginas))
    elif ext_saida == '.docx':
        _gravar_docx_texto(caminho_saida, paginas)
    elif ext_saida == '.pdf':
        _gravar_pdf(caminho_saida, paginas)
    else:
        raise ValueError("Formato de saida nao suportado: '%s'. Use .txt, .docx ou .pdf." % ext_saida)


def _gravar_pdf(caminho_saida, paginas_texto):
    doc = _fitz.open()
    largura, altura = 595, 842
    margem_x, margem_y = 45, 40
    fontsize, max_linhas = 8, 80

    for texto_pagina in paginas_texto:
        texto_pagina = _normalizar_unicode(texto_pagina)
        linhas = texto_pagina.split('\n')
        blocos = [linhas[i:i + max_linhas] for i in range(0, max(1, len(linhas)), max_linhas)]
        for bloco in blocos:
            page = doc.new_page(width=largura, height=altura)
            texto_bloco = '\n'.join(bloco)
            try:
                page.insert_text((margem_x, margem_y), texto_bloco,
                                 fontname='helv', fontsize=fontsize, color=(0, 0, 0))
            except Exception:
                texto_safe = texto_bloco.encode('latin-1', errors='replace').decode('latin-1')
                page.insert_text((margem_x, margem_y), texto_safe,
                                 fontname='helv', fontsize=fontsize, color=(0, 0, 0))

    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.pdf')
    os.close(tmp_fd)
    try:
        doc.save(tmp_path)
        doc.close()
        shutil.copy2(tmp_path, caminho_saida)
    finally:
        try: os.unlink(tmp_path)
        except OSError: pass


# ---------------------------------------------------------------------------
# Processamento principal (anonimizacao)
# ---------------------------------------------------------------------------

def processar_arquivo(caminho_entrada, caminho_saida, caminho_mapa, compliance=None):
    """Processa um arquivo com os tokens default (fluxo CLI)."""
    return aplicar_anonimizacao(caminho_entrada, caminho_saida, caminho_mapa,
                                compliance=compliance)


def registrar_compliance(rel, pasta_relatorios, nome_base):
    """Grava o relatorio individual e acrescenta ao log consolidado LGPD."""
    caminho_rel = os.path.join(pasta_relatorios, nome_base + '_relatorio.json')
    with open(caminho_rel, 'w', encoding='utf-8') as f:
        json.dump(rel, f, ensure_ascii=False, indent=2)
    log_path = os.path.join(pasta_relatorios, 'log_compliance.jsonl')
    with open(log_path, 'a', encoding='utf-8') as f:
        f.write(json.dumps(rel, ensure_ascii=False) + '\n')
    return caminho_rel


# ---------------------------------------------------------------------------
# Modulo de restauracao
# ---------------------------------------------------------------------------

def restaurar_arquivo(caminho_anonimizado, pasta_cofre, caminho_mapa=None):
    """
    Restaura um arquivo anonimizado usando o mapa de reversao do /cofre.
    Por padrao localiza o mapa pelo nome do arquivo (<nome>_mapa.json).
    caminho_mapa: permite indicar EXPLICITAMENTE outro mapa — util para
    reidentificar um documento DIFERENTE do original (ex.: parecer ou minuta
    devolvidos por uma IA) que contenha os mesmos tokens.
    Retorna (caminho_saida, trechos_revertidos).
    """
    base_nome = os.path.basename(caminho_anonimizado)
    nome_sem_sufixo = re.sub(r'_anonimizado(\.[^.]+)$', r'\1', base_nome)
    nome_base, ext = os.path.splitext(nome_sem_sufixo)
    if caminho_mapa is None:
        caminho_mapa = os.path.join(pasta_cofre, nome_base + '_mapa.json')

    if not os.path.exists(caminho_mapa):
        raise FileNotFoundError(
            'Mapa de reversao nao encontrado: %s\n'
            'O arquivo de mapa e gerado automaticamente durante a anonimizacao (v3+). '
            'Para restaurar um documento com nome diferente do original, indique o '
            'mapa explicitamente.' % caminho_mapa
        )

    with open(caminho_mapa, encoding='utf-8') as f:
        dados_mapa = json.load(f)

    mapa_reversao = dados_mapa.get('mapa_reversao', {})
    if not mapa_reversao:
        raise ValueError('Mapa de reversao vazio. Nenhuma substituicao pode ser revertida.')

    ext_lower = ext.lower()
    pasta_saida = os.path.dirname(caminho_anonimizado)
    caminho_saida = os.path.join(pasta_saida, nome_base + '_restaurado' + ext)

    if ext_lower == '.txt':
        texto = _ler_txt(caminho_anonimizado)
        mapa_resolvido = _resolver_mapa(mapa_reversao)
        revertidos = sum(texto.count(t) for t in mapa_resolvido)
        _gravar_txt(caminho_saida, reanominizar_texto(texto, mapa_reversao))

    elif ext_lower == '.docx':
        if not DOCX_DISPONIVEL:
            raise RuntimeError('python-docx nao instalado. Execute: pip install python-docx')
        revertidos = _restaurar_docx(caminho_anonimizado, caminho_saida, mapa_reversao)

    elif ext_lower == '.pdf':
        if not PDF_DISPONIVEL:
            raise RuntimeError('PyMuPDF nao instalado. Execute: pip install pymupdf')
        paginas = _ler_pdf(caminho_anonimizado)
        texto = _SEP_PAGINA.join(paginas)
        mapa_resolvido = _resolver_mapa(mapa_reversao)
        revertidos = sum(texto.count(t) for t in mapa_resolvido)
        _gravar_pdf(caminho_saida, reanominizar_texto(texto, mapa_reversao).split(_SEP_PAGINA))

    else:
        raise ValueError("Formato nao suportado para restauracao: '%s'. Use .txt, .docx ou .pdf." % ext)

    print('[OK] Arquivo restaurado: %s' % caminho_saida)
    print('     Mapa utilizado: %s' % os.path.basename(caminho_mapa))
    print('     Trechos revertidos: %d (mapa com %d chaves)' % (revertidos, len(mapa_reversao)))
    print('     Hash SHA-256 original (referencia): %s' % dados_mapa.get('hash_sha256_original', 'N/A'))
    return caminho_saida, revertidos


# ---------------------------------------------------------------------------
# Funcao principal (main)
# ---------------------------------------------------------------------------

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    pasta_entrada = os.path.join(base, 'entrada')
    pasta_saida   = os.path.join(base, 'saida')
    pasta_relat   = os.path.join(base, 'relatorios')
    pasta_cofre   = os.path.join(base, 'cofre')

    for p in (pasta_entrada, pasta_saida, pasta_relat, pasta_cofre):
        os.makedirs(p, exist_ok=True)

    # --- Modo de restauracao ---
    if len(sys.argv) >= 3 and sys.argv[1] == '--restaurar':
        caminho_alvo = sys.argv[2]
        if not os.path.isabs(caminho_alvo):
            caminho_rel = os.path.join(base, caminho_alvo)
            caminho_saida_rel = os.path.join(pasta_saida, caminho_alvo)
            if os.path.exists(caminho_rel):
                caminho_alvo = caminho_rel
            elif os.path.exists(caminho_saida_rel):
                caminho_alvo = caminho_saida_rel
        if not os.path.exists(caminho_alvo):
            print('[ERRO] Arquivo nao encontrado: %s' % caminho_alvo)
            sys.exit(1)
        # Mapa explicito opcional: python anonimizador.py --restaurar <arquivo> [mapa]
        caminho_mapa = sys.argv[3] if len(sys.argv) >= 4 else None
        if caminho_mapa and not os.path.isabs(caminho_mapa):
            candidato = os.path.join(pasta_cofre, caminho_mapa)
            caminho_mapa = candidato if os.path.exists(candidato) else os.path.join(base, caminho_mapa)
        try:
            restaurar_arquivo(caminho_alvo, pasta_cofre, caminho_mapa=caminho_mapa)
        except Exception as e:
            print('[ERRO] %s' % e)
            sys.exit(1)
        return

    # --- Formato de saida para PDFs: --saida-pdf txt|docx|pdf (padrao: pdf) ---
    formato_pdf = '.pdf'
    argv = sys.argv[1:]
    if '--saida-pdf' in argv:
        i = argv.index('--saida-pdf')
        valor = argv[i + 1].lstrip('.').lower() if i + 1 < len(argv) else ''
        if valor not in ('txt', 'docx', 'pdf'):
            print('[ERRO] --saida-pdf requer um formato: txt, docx ou pdf.')
            sys.exit(1)
        formato_pdf = '.' + valor
        if formato_pdf == '.docx' and not DOCX_DISPONIVEL:
            print('[ERRO] Saida .docx requer python-docx. Execute: pip install python-docx')
            sys.exit(1)

    # --- Modo de anonimizacao (padrao) ---
    compliance = _carregar_config(base)

    exts = {'.txt'}
    avisos = []
    if DOCX_DISPONIVEL:
        exts.add('.docx')
    else:
        avisos.append('.docx ignorado (instale python-docx)')
    if PDF_DISPONIVEL:
        exts.add('.pdf')
    else:
        avisos.append('.pdf ignorado (instale pymupdf)')
    if avisos:
        print('[AVISO] ' + ' | '.join(avisos))
        print()

    arquivos = [f for f in os.listdir(pasta_entrada)
                if os.path.splitext(f)[1].lower() in exts]

    if not arquivos:
        print('Nenhum arquivo %s encontrado em /entrada.' % ' / '.join(sorted(exts)))
        return

    log_path = os.path.join(pasta_relat, 'log_compliance.jsonl')

    print('Processando %d arquivo(s)...\n' % len(arquivos))
    for nome in arquivos:
        ce = os.path.join(pasta_entrada, nome)
        base_nome, ext = os.path.splitext(nome)
        ext_saida = formato_pdf if ext.lower() == '.pdf' else ext
        cs = os.path.join(pasta_saida, base_nome + '_anonimizado' + ext_saida)
        cm = os.path.join(pasta_cofre, base_nome + '_mapa.json')

        try:
            rel = processar_arquivo(ce, cs, cm, compliance=compliance)
        except Exception as e:
            print('[ERRO] %s: %s\n' % (nome, e))
            continue

        cr = registrar_compliance(rel, pasta_relat, base_nome)

        print('[OK] %s' % nome)
        print('     Saida    : saida/%s' % os.path.basename(cs))
        print('     Mapa     : cofre/%s  [CONFIDENCIAL]' % os.path.basename(cm))
        print('     Relatorio: relatorios/%s' % os.path.basename(cr))
        print('     Hash SHA-256: %s' % rel['hash_sha256_original'])
        print('     Total de substituicoes: %d (%d entidades unicas)'
              % (rel['total_substituicoes'], rel['entidades_unicas']))
        for tipo, qtd in sorted(rel['ocorrencias_por_tipo'].items()):
            if qtd > 0:
                print('       %s: %d' % (tipo, qtd))
        print()

    print('Log LGPD atualizado: relatorios/log_compliance.jsonl')
    print('Concluido.')


if __name__ == '__main__':
    main()
